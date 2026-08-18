"""Fase 5 — a minuta do pedido de dilação de prazo, em Word.

Duas camadas, como o resto do projeto:

- `montar(dados, analise) -> Minuta` é **puro** (testável sem `python-docx` e sem API):
  devolve a peça como lista de `(estilo, texto)`.
- `para_docx(minuta) -> bytes` é a casca fina que desenha isso num `.docx`.

A estrutura segue o template real usado pelo Jurídico
(`BKUP_TELECOM_Pedido_Dilacao_Prazo_RI_II.pdf`): endereçamento → referência → qualificação →
`I – DA TEMPESTIVIDADE` → `II – DA JUSTIFICATIVA` → `III – DO PEDIDO` → fecho.

⚠️ **Lacunas em vez de invenção.** Todo dado que o bot não tem com segurança vira
`[PREENCHER: …]` — visível no Word e localizável por busca. A entrega é uma **minuta para
revisão humana**: o Jurídico ajusta e protocola no SEI. O bot não peticiona.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date
from typing import Optional

from .dilacao import ANALISE_VAZIA, Analise

# estilos que o renderer entende
TITULO, CORPO, CENTRO, ASSINATURA = "titulo", "corpo", "centro", "assinatura"

MESES = ("janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho",
         "agosto", "setembro", "outubro", "novembro", "dezembro")

# preposições que ficam minúsculas no nome de município ("Rio das Antas", não "Rio Das Antas")
_MINUSCULAS = {"de", "da", "do", "das", "dos", "e", "d'"}

_UNIDADES = ("zero", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito",
             "nove", "dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis",
             "dezessete", "dezoito", "dezenove")
_DEZENAS = ("", "", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta",
            "oitenta", "noventa")


def lacuna(oque: str) -> str:
    return f"[PREENCHER: {oque}]"


def por_extenso(n: int) -> str:
    """Números de prazo por extenso (o template escreve '10 (dez) dias'). Cobre 0–99."""
    if n < 0 or n > 99:
        return ""
    if n < 20:
        return _UNIDADES[n]
    dez, uni = divmod(n, 10)
    return _DEZENAS[dez] + (f" e {_UNIDADES[uni]}" if uni else "")


def dias_por_extenso(n: int, unidade: str = "dias") -> str:
    """'10 (dez) dias' / '20 (vinte) dias úteis'. Dia útil ≠ dia corrido, então a unidade
    vem junto — é a mesma distinção que `processo.Prazo.unidade` preserva."""
    ext = por_extenso(n)
    u = (unidade or "dias").strip() or "dias"
    return f"{n} ({ext}) {u}" if ext else f"{n} {u}"


def cidade_formatada(municipio: str) -> str:
    """'SAO LUIS DE MONTES BELOS' → 'Sao Luis de Montes Belos'.

    O cadastro (Clientes SCM) guarda o município em caixa alta e **sem acento**; os acentos
    não são recuperáveis daqui. Esta função só conserta a caixa — a grafia acentuada vem do
    LLM quando ele responde, e este é o fallback determinístico.

    Texto que já tem minúsculas passa intacto: serve para aplicar isto também à resposta do
    LLM, que às vezes ecoa o cadastro em caixa alta ("BELO HORIZONTE") e às vezes devolve a
    grafia certa ("São Paulo") — recasear a segunda estragaria o acento da capitalização.
    """
    s = re.sub(r"\s+", " ", (municipio or "")).strip()
    if not s or any(c.islower() for c in s):
        return s
    palavras = s.lower().split(" ")
    saida = [palavras[0].capitalize()]
    saida += [p if p in _MINUSCULAS else p.capitalize() for p in palavras[1:]]
    return " ".join(saida)


_PREPOSICOES = ("no ", "na ", "nos ", "nas ", "em ", "do ", "da ", "dos ", "das ",
                "de ", "nº ", "n° ")


def com_preposicao(trecho: str) -> str:
    """Garante que o trecho encaixe depois de 'com fundamento …'.

    O artigo certo depende do substantivo ("no item 5", mas "nas orientações"), então quem
    o escolhe é o modelo — que sabe contrair em português. Esta é a rede: se ele devolver
    o trecho sem preposição nenhuma, entra um "em" neutro, que nunca erra concordância.
    (Sem isso saiu "com fundamento no orientações constantes do Manual", no ensaio real.)
    """
    s = (trecho or "").strip()
    if not s:
        return ""
    return s if s.lower().startswith(_PREPOSICOES) else f"em {s}"


def data_por_extenso(d: date) -> str:
    return f"{d.day} de {MESES[d.month - 1]} de {d.year}"


def cnpj_formatado(cnpj: str) -> str:
    s = re.sub(r"\D", "", cnpj or "")
    if len(s) != 14:
        return cnpj or ""
    return f"{s[:2]}.{s[2:5]}.{s[5:8]}/{s[8:12]}-{s[12:]}"


def certidao_de(protocolos: Optional[dict]) -> str:
    """Nº SEI da 'Certidão de Intimação Cumprida' na Lista de Protocolos — a prova da
    ciência que a seção de tempestividade cita. `protocolos` = {num: tipo}."""
    for num, tipo in (protocolos or {}).items():
        if "certid" in (tipo or "").lower() and "intima" in (tipo or "").lower():
            return str(num)
    return ""


@dataclass(frozen=True)
class Minuta:
    paragrafos: tuple[tuple[str, str], ...]
    lacunas: tuple[str, ...]
    dias_pedidos: int
    empresa: str
    processo: str
    # None = não deu para saber (sem dossiê). False = o ofício NÃO traz cláusula que
    # autorize dilação — a peça é gerada assim mesmo (decisão do usuário), mas quem revisa
    # precisa ser avisado antes de protocolar algo inadmissível. Ex.: Notificação de
    # Lançamento tributária, em que o remédio é impugnação, não dilação.
    admite_dilacao: Optional[bool] = None

    @property
    def texto(self) -> str:
        """A peça em texto puro — usada nos testes e no `--modo ensaio`."""
        return "\n\n".join(t for _, t in self.paragrafos)


def _dias_pedidos(analise: Analise, prazo_dias: Optional[int]) -> int:
    """Quantos dias pedir: o teto que o ofício admite; na falta, o prazo original.

    O prazo original é o fallback certo porque é o teto que o próprio template invoca —
    'período não superior ao originalmente concedido'. Pedir mais que o admitido seria
    convite a indeferimento.
    """
    if analise.dias > 0:
        return analise.dias
    return int(prazo_dias or 0)


def montar(dados: dict, analise: Optional[Analise] = None, *,
           hoje: Optional[date] = None) -> Minuta:
    """Monta a peça. `dados` vem da linha de `tratadas` + cadastro do cliente.

    Chaves usadas: processo, oficio_desc, doc_id, empresa, cnpj, data_ciencia, prazo_dias,
    prazo_unidade, data_limite, protocolos, cidade_cadastro, assinante, assinante_cargo,
    cidade_padrao.
    """
    # `ANALISE_VAZIA`, e não `Analise()`: só o sentinela carrega `vazia=True`, que é o que
    # distingue "o modelo leu e não achou cláusula de dilação" de "não houve leitura".
    analise = ANALISE_VAZIA if analise is None else analise
    hoje = hoje or date.today()
    lacunas: list[str] = []

    def com_lacuna(valor: str, oque: str) -> str:
        if valor:
            return valor
        lacunas.append(oque)
        return lacuna(oque)

    unidade = (dados.get("prazo_unidade") or "dias").strip() or "dias"
    prazo_dias = int(dados.get("prazo_dias") or 0)
    pedidos = _dias_pedidos(analise, prazo_dias)

    empresa = (dados.get("empresa") or "").strip()
    cnpj = cnpj_formatado(dados.get("cnpj") or "")
    processo = (dados.get("processo") or "").strip()
    ref_doc = analise.ref_documento or (dados.get("oficio_desc") or "").strip()
    doc_id = (dados.get("doc_id") or "").strip()
    # o nº SEI só é acrescentado quando a referência ainda não traz nenhum: o LLM costuma
    # devolver "Ofício 498 (SEI nº …)" pronto, e concatenar geraria "(SEI nº X) (SEI nº Y)"
    if ref_doc and doc_id and "SEI" not in ref_doc.upper():
        ref_doc = f"{ref_doc} (SEI nº {doc_id})"

    p: list[tuple[str, str]] = []

    # --- endereçamento -------------------------------------------------------------
    p.append((TITULO, "À AGÊNCIA NACIONAL DE TELECOMUNICAÇÕES – ANATEL"))
    if analise.orgao:
        p += [(CORPO, linha) for linha in analise.orgao]
    else:
        lacunas.append("unidade da Anatel (Gerência/Coordenação)")
        p.append((CORPO, lacuna("unidade da Anatel (Gerência/Coordenação)")))

    # --- referência ----------------------------------------------------------------
    p.append((TITULO, f"Processo nº {com_lacuna(processo, 'número do processo')}"))
    p.append((CORPO, f"Ref.: {com_lacuna(ref_doc, 'documento que fixou o prazo')}"))
    p.append((TITULO, "Assunto: Pedido de dilação de prazo"))

    # --- qualificação --------------------------------------------------------------
    fundamento = (f"com fundamento {com_preposicao(analise.fundamento)}"
                  if analise.fundamento
                  else com_lacuna("", "fundamento que autoriza a dilação (item do ofício/RI)"))
    # "e em X" (não "e nas disposições do X"): a norma tanto pode ser "Lei nº 9.998…" quanto
    # "Regulamento de…", e um artigo fixo erraria a concordância em metade dos casos.
    # Norma já contida no fundamento não é repetida — o modelo às vezes devolve as duas iguais.
    norma = ""
    if analise.norma and analise.norma not in analise.fundamento:
        norma = f" e {com_preposicao(analise.norma)}"
    p.append((CORPO,
              f"{com_lacuna(empresa, 'razão social da requerente').upper()}, pessoa jurídica "
              f"de direito privado, inscrita no CNPJ sob o nº "
              f"{com_lacuna(cnpj, 'CNPJ da requerente')}, já qualificada nos autos do processo "
              f"em epígrafe, vem, respeitosamente, por meio de seu procurador que esta "
              f"subscreve, {fundamento}{norma}, apresentar PEDIDO DE DILAÇÃO DE PRAZO, pelas "
              f"razões de fato e de direito a seguir expostas."))

    # --- I — tempestividade --------------------------------------------------------
    p.append((TITULO, "I – DA TEMPESTIVIDADE"))
    ciencia = com_lacuna((dados.get("data_ciencia") or "").strip(), "data da ciência")
    certidao = certidao_de(dados.get("protocolos"))
    prova = (f", conforme Certidão de Intimação Cumprida (SEI nº {certidao})"
             if certidao else "")
    prazo_txt = (dias_por_extenso(prazo_dias, unidade) if prazo_dias
                 else com_lacuna("", "prazo original concedido"))
    p.append((CORPO,
              f"A intimação foi cumprida por consulta direta em {ciencia}{prova}, tendo sido "
              f"fixado o prazo de {prazo_txt} para atendimento ao que foi requerido, contado a "
              f"partir do primeiro dia útil seguinte ao recebimento."))
    limite = (dados.get("data_limite") or "").strip()
    vencimento = f", cujo termo final é {limite}" if limite else ""
    p.append((CORPO,
              f"O presente pedido é protocolado ainda dentro do prazo originalmente "
              f"concedido{vencimento}, em estrita observância ao que admite o próprio "
              f"instrumento de intimação. A Requerente está ciente de que o pedido de dilação "
              f"não suspende nem interrompe o prazo originalmente concedido."))

    # --- II — justificativa ---------------------------------------------------------
    p.append((TITULO, "II – DA JUSTIFICATIVA"))
    if analise.justificativa:
        p += [(CORPO, par) for par in analise.justificativa]
    else:
        lacunas.append("justificativa do pedido")
        p.append((CORPO, lacuna("justificativa do pedido — descrever, com base no que o "
                                "ofício exige, por que o prazo original é insuficiente")))
    p.append((CORPO,
              "A presente solicitação não tem, portanto, qualquer propósito protelatório: visa "
              "exclusivamente assegurar a entrega de resposta íntegra, organizada e fidedigna, "
              "em prestígio aos princípios da verdade material e da eficiência que regem o "
              "processo administrativo (art. 2º da Lei nº 9.784/1999), e em colaboração com a "
              "atividade fiscalizatória em curso."))

    # --- III — pedido ---------------------------------------------------------------
    p.append((TITULO, "III – DO PEDIDO"))
    quanto = (dias_por_extenso(pedidos, unidade) if pedidos
              else com_lacuna("", "quantidade de dias de dilação a requerer"))
    p.append((CORPO,
              f"Ante o exposto, requer-se a dilação do prazo para atendimento ao "
              f"{ref_doc or 'documento em referência'} por {quanto}, período não superior ao "
              f"originalmente concedido, contando-se o novo prazo, em caso de deferimento, a "
              f"partir do primeiro dia útil posterior ao vencimento do prazo original."))
    p.append((CORPO,
              "A Requerente reitera seu integral compromisso de colaboração com a fiscalização "
              "e informa que, independentemente do deferimento, prosseguirá na consolidação "
              "imediata da documentação, podendo, inclusive, antecipar a entrega tão logo "
              "concluída."))

    # --- fecho ----------------------------------------------------------------------
    p.append((CENTRO, "Nestes termos, pede deferimento."))
    cidade = cidade_formatada(analise.cidade or dados.get("cidade_cadastro") or "") \
        or (dados.get("cidade_padrao") or "").strip()
    if not cidade:
        lacunas.append("cidade do fecho")
        cidade = lacuna("cidade")
    p.append((CENTRO, f"{cidade}, {data_por_extenso(hoje)}."))
    p.append((ASSINATURA, com_lacuna((dados.get("assinante") or "").strip(), "assinante")))
    cargo = (dados.get("assinante_cargo") or "").strip()
    if cargo:
        p.append((ASSINATURA, cargo))

    return Minuta(paragrafos=tuple(p), lacunas=tuple(lacunas), dias_pedidos=pedidos,
                  empresa=empresa, processo=processo,
                  admite_dilacao=(None if analise.vazia else analise.admite))


IDIOMA = "pt-BR"


def _definir_idioma(doc, idioma: str = IDIOMA) -> None:
    """Marca o documento como português do Brasil.

    O `.docx` que o python-docx cria herda o idioma do template padrão dele (inglês), então
    o Word marca **cada palavra em português** como erro de ortografia — a peça chega ao
    Jurídico coberta de vermelho. Isto declara o idioma no estilo `Normal`, de onde todo o
    resto herda, e o corretor passa a revisar em pt-BR (que é o que se quer numa minuta:
    corretor ligado, no idioma certo).
    """
    from docx.oxml.ns import qn

    def marcar(rpr) -> None:
        for velho in rpr.findall(qn("w:lang")):  # substitui, não acumula
            rpr.remove(velho)
        lang = rpr.makeelement(qn("w:lang"), {})
        lang.set(qn("w:val"), idioma)            # texto latino
        lang.set(qn("w:eastAsia"), idioma)
        rpr.append(lang)

    marcar(doc.styles["Normal"].element.get_or_add_rPr())
    # e também na raiz (`docDefaults`), que no template do python-docx vem "en-US": o estilo
    # Normal já cobriria por herança, mas deixar a raiz em inglês faz qualquer texto que não
    # herde de Normal voltar a ser corrigido em inglês.
    raiz = doc.styles.element.find(qn("w:docDefaults"))
    if raiz is not None:
        rpr_default = raiz.find(qn("w:rPrDefault"))
        if rpr_default is not None:
            rpr = rpr_default.find(qn("w:rPr"))
            if rpr is None:
                rpr = rpr_default.makeelement(qn("w:rPr"), {})
                rpr_default.append(rpr)
            marcar(rpr)


def para_docx(minuta: Minuta) -> bytes:
    """Renderiza o `.docx`. Casca fina — toda a decisão de conteúdo está em `montar`."""
    import io

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()
    estilo = doc.styles["Normal"]
    estilo.font.name = "Times New Roman"
    estilo.font.size = Pt(12)
    _definir_idioma(doc)
    for secao in doc.sections:
        secao.top_margin = secao.bottom_margin = Cm(2.5)
        secao.left_margin = secao.right_margin = Cm(2.5)

    alinhamento = {
        TITULO: WD_ALIGN_PARAGRAPH.JUSTIFY,
        CORPO: WD_ALIGN_PARAGRAPH.JUSTIFY,
        CENTRO: WD_ALIGN_PARAGRAPH.CENTER,
        ASSINATURA: WD_ALIGN_PARAGRAPH.CENTER,
    }
    for tipo, texto in minuta.paragrafos:
        par = doc.add_paragraph()
        par.alignment = alinhamento[tipo]
        par.paragraph_format.space_after = Pt(10)
        par.paragraph_format.line_spacing = 1.5
        run = par.add_run(texto)
        run.bold = tipo in (TITULO, ASSINATURA)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def nome_arquivo(minuta: Minuta, doc_id: str = "") -> str:
    """Nome do arquivo publicado. Sanitizado para caminho de SharePoint."""
    empresa = re.sub(r"[\\/:*?\"<>|%#]+", "-", minuta.empresa or "").strip() or "Requerente"
    sufixo = f" ({doc_id})" if doc_id else ""
    return f"Pedido de Dilacao de Prazo - {empresa}{sufixo}.docx"
